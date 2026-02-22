"""
CoRe — Registo de Consulta Cardio-Renal
Extração automática de dados clínicos via Gemini → Google Sheets
"""

import streamlit as st
from google import genai
from google.genai import types as genai_types
import gspread
from google.oauth2.service_account import Credentials
import json
import re
from datetime import datetime, date
from docx import Document
import pdfplumber
import io
import pandas as pd

# ─── PAGE CONFIG ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="CoRe — Registo Clínico",
    page_icon="🫀",
    layout="wide"
)

# ─── CONSTANTS ────────────────────────────────────────────────────────────────
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]
SHEET_DOENTES = "Doentes"
SHEET_VISITAS = "Visitas_Análises"
SHEET_EVENTOS = "Eventos"

HEADERS_DOENTES = [
    "N_Processo", "Data_Nascimento", "Idade", "Sexo", "Localidade",
    "Profissao", "Frailty_CFS", "Referenciacao",
    # FRCV
    "DM2", "Tabagismo", "HTA", "Dislipidemia", "Obesidade",
    "SAOS", "Sedentarismo", "HxFamiliar_DC",
    # Comorbilidades
    "DAP", "DPOC", "Doenca_hepatica", "HBP", "FA", "Outras_comorbilidades",
    # IC
    "IC_FE_tipo", "IC_Etiologia", "IC_FEVE_atual_pct", "IC_FEVE_trajetoria",
    # DRC
    "DRC_Grau", "DRC_Albuminuria", "DRC_Etiologia",
    # Congestão + POCUS
    "Fenotipo_congestao",
    "POCUS_FE_pct", "POCUS_EE_ratio", "POCUS_LinhasB_N", "POCUS_VCI_mm",
    # Medicação (12 classes × 3 colunas)
    "RASi", "RASi_farmaco", "RASi_dose",
    "MRA", "MRA_farmaco", "MRA_dose",
    "iSGLT2", "iSGLT2_farmaco", "iSGLT2_dose",
    "GLP1RA", "GLP1RA_farmaco", "GLP1RA_dose",
    "Estatina", "Estatina_farmaco", "Estatina_dose",
    "Diuretico_ansa", "Diuretico_ansa_farmaco", "Diuretico_ansa_dose",
    "Diuretico_tiazida", "Diuretico_tiazida_farmaco", "Diuretico_tiazida_dose",
    "Acetazolamida", "Acetazolamida_farmaco", "Acetazolamida_dose",
    "BetaBloqueante", "BetaBloqueante_farmaco", "BetaBloqueante_dose",
    "Antiagregante", "Antiagregante_farmaco", "Antiagregante_dose",
    "Anticoagulante", "Anticoagulante_farmaco", "Anticoagulante_dose",
    "Ivabradina", "Ivabradina_farmaco", "Ivabradina_dose",
    "Data_ultima_consulta"
]

HEADERS_VISITAS = [
    "N_Processo", "Data_consulta",
    # Função renal
    "Ureia", "Creatinina", "Cistatina_C", "TFGe_CKD_EPI_CrCist",
    "RACu_mg_g", "RPC_mg_g", "Na_urinario",
    # Proteínas / Hepático
    "Albumina", "ALT", "AST", "GGT", "Bilirrubina_total",
    # Eletrólitos / Minerais
    "Na", "K", "Cl", "Ca", "P", "Mg",
    # Endócrino
    "PTH", "Vit_D",
    # Biomarcadores
    "NT_proBNP", "BNP", "CA125",
    # Hemograma (selecionado)
    "Hgb", "Leucocitos", "Plaquetas",
    # Gasimetria (selecionado)
    "HCO3", "Ca_ionizado",
    # Urina
    "Sumario_urina",
    # Sintomas
    "NYHA", "CCS", "Ortopneia", "Bendopneia", "Edemas_MI",
    "Claudicacao_intermitente", "Palpitacoes",
    # Exame físico
    "Peso_kg", "Altura_m", "IMC", "TA_sist", "TA_diast", "FC", "SpO2"
]

HEADERS_EVENTOS = [
    "N_Processo", "Data_evento", "Tipo_evento", "Causa_descricao", "Data_registo"
]

MED_LABELS = {
    "rasi":             "RASi (IECA/ARA/ARNi)",
    "mra":              "MRA",
    "isglt2":           "iSGLT2",
    "glp1ra":           "GLP-1RA",
    "estatina":         "Estatina",
    "diuretico_ansa":   "Diurético de ansa",
    "diuretico_tiazida":"Diurético tiazida",
    "acetazolamida":    "Acetazolamida",
    "beta_bloqueante":  "Beta-bloqueante",
    "antiagregante":    "Antiagregante",
    "anticoagulante":   "Anticoagulante",
    "ivabradina":       "Ivabradina",
}

# ─── AUTENTICAÇÃO SIMPLES ─────────────────────────────────────────────────────
def check_password() -> bool:
    """Verifica palavra-passe simples. Devolve True se autenticado."""
    if st.session_state.get("authenticated"):
        return True
    st.title("🫀 CoRe — Registo Cardio-Renal")
    pwd = st.text_input("Palavra-passe de acesso", type="password")
    if st.button("Entrar"):
        if pwd == st.secrets.get("app_password", ""):
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("Palavra-passe incorrecta.")
    return False

# ─── GOOGLE SHEETS ────────────────────────────────────────────────────────────
@st.cache_resource
def get_gspread_client():
    creds = Credentials.from_service_account_info(
        dict(st.secrets["gcp_service_account"]), scopes=SCOPES
    )
    return gspread.authorize(creds)

def get_spreadsheet():
    client = get_gspread_client()
    return client.open_by_key(st.secrets["spreadsheet_id"])

def get_or_create_sheet(spreadsheet, name: str, headers: list):
    try:
        ws = spreadsheet.worksheet(name)
    except gspread.WorksheetNotFound:
        ws = spreadsheet.add_worksheet(title=name, rows=2000, cols=len(headers) + 5)
        ws.append_row(headers)
    return ws

def update_doentes_sheet(ws, n_processo: str, row: list):
    """Atualiza linha existente ou acrescenta nova."""
    col_a = ws.col_values(1)          # lista, 0-indexed; posição 0 = cabeçalho
    if n_processo in col_a:
        row_idx = col_a.index(n_processo) + 1   # gspread é 1-indexed
        ws.delete_rows(row_idx)
        ws.insert_row(row, row_idx)
    else:
        ws.append_row(row)

# ─── PARSERS DE FICHEIRO ──────────────────────────────────────────────────────
def parse_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    # Incluir texto em tabelas, se existirem
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)

def parse_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts).strip()

# ─── PROMPT & EXTRAÇÃO LLM ───────────────────────────────────────────────────
EXTRACTION_PROMPT = """
És um assistente especializado em extração de dados clínicos de consultas de Cardiologia-Nefrologia (síndrome cardiorrenal).
Analisa o registo clínico em Português e devolve APENAS um objeto JSON válido com a estrutura indicada abaixo.

REGRAS DE EXTRAÇÃO:
- Booleanos: true ou false (nunca "Sim"/"Não")
- Números: só o valor numérico, sem unidades
- Datas: formato "YYYY-MM-DD"  |  null se ausente
- Strings não encontradas: null
- Sexo: "M" ou "F"
- NYHA: inteiro 1–4  |  CCS: inteiro 0–4  |  Frailty CFS: inteiro 1–9
- IC tipo FE: "FEr" (<40%), "FEp" (≥50%), "FEmr" (40–49%), ou null
- DRC Grau: "G1"/"G2"/"G3a"/"G3b"/"G4"/"G5" ou null
- DRC Albuminúria: "A1" (<30 mg/g), "A2" (30–300 mg/g), "A3" (>300 mg/g) ou null
- Fenótipo congestão: "Tecidular"/"Vascular"/"Misto"/"Ausente" ou null
- Referenciação: "Cardio"/"Nefro"/"Outra"/"Pós-internamento" ou null
- TFGe: valor numérico (preferencialmente CKD-EPI Cr-Cist se disponível)
- RACu/RAC: valor em mg/g como número
- NT-proBNP: valor em pg/mL como número
- Linhas B: número de campos pulmonares com linhas B (ex: "7/8" → 7)
- VCI: diâmetro em mm como número
- FE: percentagem como número (ex: "40%" → 40)
- Para medicação: "presente" = true se o fármaco constar na medicação actual do doente
- RASi inclui: IECA (lisinopril, ramipril, enalapril, perindopril...), ARA (valsartan, losartan, olmesartan...), ARNi (sacubitril/valsartan = Entresto)
- MRA inclui: espironolactona, finerenona, eplerenona
- iSGLT2 inclui: dapagliflozina, empagliflozina, canagliflozina
- GLP-1RA inclui: semaglutido, liraglutido, dulaglutido, exenatido
- Antiagregante inclui: AAS/ácido acetilsalicílico, clopidogrel, ticagrelor, prasugrel
- Anticoagulante inclui: apixabano, rivaroxabano, dabigatrano, edoxabano, varfarina
- Beta-bloqueante inclui: carvedilol, bisoprolol, nebivolol, metoprolol, atenolon

TEXTO CLÍNICO:
{texto}

Responde EXCLUSIVAMENTE com o JSON abaixo preenchido (sem markdown, sem texto extra):

{
  "doente": {
    "data_nascimento": null,
    "sexo": null,
    "localidade": null,
    "profissao": null,
    "frailty_cfs": null,
    "referenciacao": null,
    "frcv": {
      "dm2": null,
      "tabagismo": null,
      "hta": null,
      "dislipidemia": null,
      "obesidade": null,
      "saos": null,
      "sedentarismo": null,
      "hx_familiar_dc": null
    },
    "comorbilidades": {
      "dap": null,
      "dpoc": null,
      "doenca_hepatica": null,
      "hbp": null,
      "fa": null,
      "outras": null
    },
    "ic": {
      "tipo_fe": null,
      "etiologia": null,
      "feve_atual": null,
      "feve_trajetoria": null
    },
    "drc": {
      "grau": null,
      "albuminuria": null,
      "etiologia": null
    },
    "fenotipo_congestao": null,
    "pocus": {
      "fe_pct": null,
      "ee_ratio": null,
      "linhas_b_n": null,
      "vci_mm": null
    },
    "medicacao": {
      "rasi":              {"presente": null, "farmaco": null, "dose": null},
      "mra":               {"presente": null, "farmaco": null, "dose": null},
      "isglt2":            {"presente": null, "farmaco": null, "dose": null},
      "glp1ra":            {"presente": null, "farmaco": null, "dose": null},
      "estatina":          {"presente": null, "farmaco": null, "dose": null},
      "diuretico_ansa":    {"presente": null, "farmaco": null, "dose": null},
      "diuretico_tiazida": {"presente": null, "farmaco": null, "dose": null},
      "acetazolamida":     {"presente": null, "farmaco": null, "dose": null},
      "beta_bloqueante":   {"presente": null, "farmaco": null, "dose": null},
      "antiagregante":     {"presente": null, "farmaco": null, "dose": null},
      "anticoagulante":    {"presente": null, "farmaco": null, "dose": null},
      "ivabradina":        {"presente": null, "farmaco": null, "dose": null}
    }
  },
  "visita": {
    "data_consulta": null,
    "analises": {
      "ureia": null,
      "creatinina": null,
      "cistatina_c": null,
      "tfge_ckd_epi_crcist": null,
      "racu": null,
      "rpc": null,
      "na_urinario": null,
      "albumina": null,
      "alt": null,
      "ast": null,
      "ggt": null,
      "bilirrubina_total": null,
      "na": null,
      "k": null,
      "cl": null,
      "ca": null,
      "p": null,
      "mg": null,
      "pth": null,
      "vit_d": null,
      "nt_probnp": null,
      "bnp": null,
      "ca125": null,
      "hgb": null,
      "leucocitos": null,
      "plaquetas": null,
      "hco3": null,
      "ca_ionizado": null,
      "sumario_urina": null
    },
    "sintomas": {
      "nyha": null,
      "ccs": null,
      "ortopneia": null,
      "bendopneia": null,
      "edemas_mi": null,
      "claudicacao_intermitente": null,
      "palpitacoes": null
    },
    "exame_fisico": {
      "peso_kg": null,
      "altura_m": null,
      "imc": null,
      "ta_sist": null,
      "ta_diast": null,
      "fc": null,
      "spo2": null
    }
  }
}
"""
def extract_with_gemini(texto: str) -> dict:
    client = genai.Client(api_key=st.secrets["gemini_api_key"])
    prompt = EXTRACTION_PROMPT.replace("{texto}", texto)
    response = client.models.generate_content(
        model="gemini-2.0-flash-001",
        contents=prompt,
        config=genai_types.GenerateContentConfig(
            temperature=0.1,
            response_mime_type="application/json",
        )
    )
    raw = response.text.strip()
    # Limpar eventual markdown
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)

# ─── HELPERS DE VALOR ──────────────────────────────────────────────────────────
def sv(val) -> str:
    """safe value → string"""
    if val is None:
        return ""
    if val is True:
        return "Sim"
    if val is False:
        return "Não"
    return str(val)

def calculate_age(dob_str: str):
    try:
        dob = datetime.strptime(dob_str, "%Y-%m-%d").date()
        today = date.today()
        return today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
    except Exception:
        return None

# ─── CONSTRUTORES DE LINHAS PARA O SHEET ─────────────────────────────────────
def build_doentes_row(n_processo: str, extracted: dict) -> list:
    d  = extracted["doente"]
    frcv = d.get("frcv", {})
    co   = d.get("comorbilidades", {})
    ic   = d.get("ic", {})
    drc  = d.get("drc", {})
    poc  = d.get("pocus", {})
    med  = d.get("medicacao", {})
    age  = calculate_age(d.get("data_nascimento") or "")

    def med3(key):
        m = med.get(key) or {}
        return [sv(m.get("presente")), sv(m.get("farmaco")), sv(m.get("dose"))]

    return [
        n_processo,
        sv(d.get("data_nascimento")), sv(age),
        sv(d.get("sexo")), sv(d.get("localidade")),
        sv(d.get("profissao")), sv(d.get("frailty_cfs")),
        sv(d.get("referenciacao")),
        # FRCV
        sv(frcv.get("dm2")), sv(frcv.get("tabagismo")), sv(frcv.get("hta")),
        sv(frcv.get("dislipidemia")), sv(frcv.get("obesidade")),
        sv(frcv.get("saos")), sv(frcv.get("sedentarismo")), sv(frcv.get("hx_familiar_dc")),
        # Comorbilidades
        sv(co.get("dap")), sv(co.get("dpoc")), sv(co.get("doenca_hepatica")),
        sv(co.get("hbp")), sv(co.get("fa")), sv(co.get("outras")),
        # IC
        sv(ic.get("tipo_fe")), sv(ic.get("etiologia")),
        sv(ic.get("feve_atual")), sv(ic.get("feve_trajetoria")),
        # DRC
        sv(drc.get("grau")), sv(drc.get("albuminuria")), sv(drc.get("etiologia")),
        # Congestão + POCUS
        sv(d.get("fenotipo_congestao")),
        sv(poc.get("fe_pct")), sv(poc.get("ee_ratio")),
        sv(poc.get("linhas_b_n")), sv(poc.get("vci_mm")),
        # Medicação (12 classes × 3)
        *med3("rasi"), *med3("mra"), *med3("isglt2"), *med3("glp1ra"),
        *med3("estatina"), *med3("diuretico_ansa"), *med3("diuretico_tiazida"),
        *med3("acetazolamida"), *med3("beta_bloqueante"),
        *med3("antiagregante"), *med3("anticoagulante"), *med3("ivabradina"),
        # Data
        date.today().isoformat(),
    ]

def build_visitas_row(n_processo: str, extracted: dict) -> list:
    v  = extracted["visita"]
    a  = v.get("analises", {})
    s  = v.get("sintomas", {})
    ef = v.get("exame_fisico", {})
    return [
        n_processo, sv(v.get("data_consulta")),
        sv(a.get("ureia")), sv(a.get("creatinina")), sv(a.get("cistatina_c")),
        sv(a.get("tfge_ckd_epi_crcist")),
        sv(a.get("racu")), sv(a.get("rpc")), sv(a.get("na_urinario")),
        sv(a.get("albumina")),
        sv(a.get("alt")), sv(a.get("ast")), sv(a.get("ggt")), sv(a.get("bilirrubina_total")),
        sv(a.get("na")), sv(a.get("k")), sv(a.get("cl")), sv(a.get("ca")),
        sv(a.get("p")), sv(a.get("mg")),
        sv(a.get("pth")), sv(a.get("vit_d")),
        sv(a.get("nt_probnp")), sv(a.get("bnp")), sv(a.get("ca125")),
        sv(a.get("hgb")), sv(a.get("leucocitos")), sv(a.get("plaquetas")),
        sv(a.get("hco3")), sv(a.get("ca_ionizado")),
        sv(a.get("sumario_urina")),
        # Sintomas
        sv(s.get("nyha")), sv(s.get("ccs")),
        sv(s.get("ortopneia")), sv(s.get("bendopneia")), sv(s.get("edemas_mi")),
        sv(s.get("claudicacao_intermitente")), sv(s.get("palpitacoes")),
        # Exame físico
        sv(ef.get("peso_kg")), sv(ef.get("altura_m")), sv(ef.get("imc")),
        sv(ef.get("ta_sist")), sv(ef.get("ta_diast")), sv(ef.get("fc")), sv(ef.get("spo2")),
    ]

# ─── COMPONENTES UI ───────────────────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.header("ℹ️ Instruções")
        st.markdown("""
**1.** Introduz o N° de processo
**2.** Faz upload da nota de consulta (`.docx`)
**3.** *(Opcional)* Faz upload das análises (`.pdf`)
**4.** Clica **Processar Consulta**
**5.** Revê os dados extraídos
**6.** Clica **Guardar no Google Sheets**

---
📋 **Folha Doentes** → dados mais recentes (1 linha/doente)
🧪 **Folha Visitas** → histórico analítico (1 linha/consulta)
🏥 **Folha Eventos** → preenchimento manual
        """)
        st.divider()
        sheet_id = st.secrets.get("spreadsheet_id", "")
        if sheet_id:
            url = f"https://docs.google.com/spreadsheets/d/{sheet_id}"
            st.link_button("📊 Abrir Google Sheet", url)

def render_review(extracted: dict):
    """Mostra resumo dos dados extraídos para revisão."""
    d   = extracted["doente"]
    v   = extracted["visita"]
    ic  = d.get("ic", {})
    drc = d.get("drc", {})
    poc = d.get("pocus", {})
    med = d.get("medicacao", {})
    a   = v.get("analises", {})
    s   = v.get("sintomas", {})
    ef  = v.get("exame_fisico", {})

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 👤 Identificação")
        st.write(f"**Data nasc.:** {sv(d.get('data_nascimento')) or '—'} &nbsp;|&nbsp; **Sexo:** {sv(d.get('sexo')) or '—'}")
        st.write(f"**Localidade:** {sv(d.get('localidade')) or '—'} &nbsp;|&nbsp; **Profissão:** {sv(d.get('profissao')) or '—'}")
        st.write(f"**Referenciação:** {sv(d.get('referenciacao')) or '—'} &nbsp;|&nbsp; **Frailty CFS:** {sv(d.get('frailty_cfs')) or '—'}")

        st.markdown("#### 🫀 Contexto Cardio-Renal")
        st.write(f"**IC:** {sv(ic.get('tipo_fe')) or '—'} — {sv(ic.get('etiologia')) or '—'} | FEVE {sv(ic.get('feve_atual')) or '—'}% | {sv(ic.get('feve_trajetoria')) or '—'}")
        st.write(f"**DRC:** {sv(drc.get('grau')) or '—'} {sv(drc.get('albuminuria')) or '—'} — {sv(drc.get('etiologia')) or '—'}")
        st.write(f"**Congestão:** {sv(d.get('fenotipo_congestao')) or '—'}")

        st.markdown("#### 🔬 POCUS")
        st.write(f"FE **{sv(poc.get('fe_pct')) or '—'}%** | E/E' **{sv(poc.get('ee_ratio')) or '—'}** | Linhas B **{sv(poc.get('linhas_b_n')) or '—'}** campos | VCI **{sv(poc.get('vci_mm')) or '—'} mm**")

        st.markdown("#### 🩺 Exame físico & Sintomas")
        st.write(f"**NYHA:** {sv(s.get('nyha')) or '—'} | **CCS:** {sv(s.get('ccs')) or '—'}")
        st.write(f"**Ortopneia:** {sv(s.get('ortopneia')) or '—'} | **Bendopneia:** {sv(s.get('bendopneia')) or '—'} | **Edemas MI:** {sv(s.get('edemas_mi')) or '—'}")
        st.write(f"**Peso:** {sv(ef.get('peso_kg')) or '—'} kg | **IMC:** {sv(ef.get('imc')) or '—'}")
        st.write(f"**TA:** {sv(ef.get('ta_sist')) or '—'}/{sv(ef.get('ta_diast')) or '—'} mmHg | **FC:** {sv(ef.get('fc')) or '—'} bpm | **SpO₂:** {sv(ef.get('spo2')) or '—'}%")

    with col2:
        st.markdown("#### 💊 Medicação")
        for key, label in MED_LABELS.items():
            m = med.get(key) or {}
            if m.get("presente") is True:
                farmaco = sv(m.get("farmaco")) or "—"
                dose    = sv(m.get("dose")) or ""
                st.write(f"✅ **{label}:** {farmaco} {dose}".strip())
            elif m.get("presente") is False:
                st.write(f"❌ **{label}**")
            else:
                st.write(f"❓ **{label}:** não identificado")

        st.markdown("#### 🧪 Análises (principais)")
        def lab(label, val, unit=""):
            v_str = sv(val) or "—"
            st.write(f"**{label}:** {v_str}{' ' + unit if v_str != '—' and unit else ''}")

        lab("TFGe (CKD-EPI Cr-Cist)", a.get("tfge_ckd_epi_crcist"), "mL/min")
        lab("Creatinina", a.get("creatinina"), "mg/dL")
        lab("Cistatina C", a.get("cistatina_c"), "mg/L")
        lab("RACu", a.get("racu"), "mg/g")
        lab("NT-proBNP", a.get("nt_probnp"), "pg/mL")
        lab("K", a.get("k"), "mEq/L")
        lab("Na", a.get("na"), "mEq/L")
        lab("Hgb", a.get("hgb"), "g/dL")
        lab("Albumina", a.get("albumina"), "g/dL")
        lab("HCO₃⁻", a.get("hco3"), "mmol/L")

# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    if not check_password():
        return

    render_sidebar()

    st.title("🫀 CoRe — Registo de Consulta Cardio-Renal")
    st.caption("Upload da nota de consulta → extração automática por IA → Google Sheets")

    # Inicializar session state
    if "ready_to_save" not in st.session_state:
        st.session_state["ready_to_save"] = False

    # ── INPUTS ────────────────────────────────────────────────────────────────
    n_processo = st.text_input(
        "N° de Processo *",
        placeholder="Ex: 123456",
        help="Introduz manualmente — não é extraído do documento"
    )

    col_up1, col_up2 = st.columns(2)
    with col_up1:
        docx_file = st.file_uploader(
            "📋 Nota de consulta (.docx) *", type=["docx"]
        )
    with col_up2:
        pdf_file = st.file_uploader(
            "🧪 Análises laboratoriais (.pdf) — opcional", type=["pdf"]
        )

    # ── PROCESSARB────────────────────────────────────────────────────────────
    can_process = bool(n_processo and docx_file)
    if st.button("⚡ Processar Consulta", type="primary", disabled=not can_process):
        texto_total = ""

        with st.spinner("A extrair texto dos ficheiros…"):
            try:
                texto_total += parse_docx(docx_file.read())
            except Exception as e:
                st.error(f"Erro a ler o .docx: {e}")
                return
            if pdf_file:
                try:
                    texto_total += "\n\n=== ANÁLISES LABORATORIAIS ===\n"
                    texto_total += parse_pdf(pdf_file.read())
                except Exception as e:
                    st.warning(f"Não foi possível ler o PDF das análises: {e}")

        with st.spinner("A enviar para o Gemini e a extrair dados estruturados…"):
            try:
                extracted = extract_with_gemini(texto_total)
                st.session_state["extracted"]   = extracted
                st.session_state["n_processo"]  = n_processo
                st.session_state["ready_to_save"] = True
                st.success("✅ Extração concluída! Revê os dados abaixo antes de guardar.")
            except json.JSONDecodeError as e:
                st.error(f"O Gemini devolveu uma resposta que não é JSON válido: {e}")
                return
            except Exception as e:
                st.error(f"Erro na extração: {e}")
                return

    # ── REVISÃO & GUARDAR ────────────────────────────────────────────────────
    if st.session_state.get("ready_to_save") and "extracted" in st.session_state:
        extracted  = st.session_state["extracted"]
        n_processo = st.session_state["n_processo"]

        st.divider()
        st.subheader("📋 Revisão dos dados extraídos")
        st.caption("Verifica antes de guardar. Podes corrigir directamente no Google Sheet após guardar.")
        render_review(extracted)

        # JSON bruto (debug)
        with st.expander("🔍 Ver JSON completo (debug)"):
            st.json(extracted)

        st.divider()
        col_ok, col_cancel, _ = st.columns([1, 1, 4])

        with col_ok:
            if st.button("💾 Guardar no Google Sheets", type="primary"):
                with st.spinner("A guardar…"):
                    try:
                        ss = get_spreadsheet()
                        ws_d = get_or_create_sheet(ss, SHEET_DOENTES, HEADERS_DOENTES)
                        ws_v = get_or_create_sheet(ss, SHEET_VISITAS, HEADERS_VISITAS)
                        ws_e = get_or_create_sheet(ss, SHEET_EVENTOS, HEADERS_EVENTOS)

                        update_doentes_sheet(ws_d, n_processo, build_doentes_row(n_processo, extracted))
                        ws_v.append_row(build_visitas_row(n_processo, extracted))

                        st.success(f"✅ Dados do processo **{n_processo}** guardados com sucesso!")
                        st.balloons()
                        st.session_state["ready_to_save"] = False
                    except Exception as e:
                        st.error(f"Erro ao guardar no Google Sheets: {e}")

        with col_cancel:
            if st.button("🗑️ Cancelar"):
                st.session_state["ready_to_save"] = False
                st.session_state.pop("extracted", None)
                st.rerun()


if __name__ == "__main__":
    main()
